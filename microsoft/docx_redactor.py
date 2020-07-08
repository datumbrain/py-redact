import sys
import os.path
import re

from docx import Document

class DocxRedactor:
    def __init__(self, doc_obj_path, regexes, replace_char):
        self.doc_obj_path = doc_obj_path
        self.regexes = regexes
        self.replace_char = replace_char
    
    def __redact_helper__(self,doc_obj):
        for reg in self.regexes:
            regex = re.compile(reg)
            for p in doc_obj.paragraphs:
                if regex.search(p.text):
                    inline = p.runs
                    for i in range(len(inline)):
                        if regex.search(inline[i].text):
                            text = regex.sub(self.replace_char * len(inline[i].text), inline[i].text)
                            inline[i].text = text
            for table in doc_obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self.__redact_helper__(cell)
        

    def redact(self, output_file_path):
        doc_obj = Document(self.doc_obj_path)
        self.__redact_helper__(doc_obj)
        doc_obj.save(output_file_path)
        if(self.doc_obj_path == output_file_path):
            print("Warning: Input and Output files are same!")
        print("Updated file saved as: " + output_file_path)