import sys
import os.path
import re
from docx import Document

class DocxRedactor:
    def __init__(self, doc_obj_path, regexes, replace_char):
        self.doc_obj_path = doc_obj_path
        self.regexes = regexes
        self.replace_char = replace_char
    
    def __redact_helper__(self,doc_obj):
        for reg in self.regexes:
            regex = re.compile(reg)
            for p in doc_obj.paragraphs:
                if regex.search(p.text):
                    inline = p.runs
                    for i in range(len(inline)):
                        if regex.search(inline[i].text):
                            text = regex.sub(self.replace_char * len(inline[i].text), inline[i].text)
                            inline[i].text = text
            for table in doc_obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self.__redact_helper__(cell)
        

    def redact(self, output_file_path):
        doc_obj = Document(self.doc_obj_path)
        self.__redact_helper__(doc_obj)
        doc_obj.save(output_file_path)
        if(self.doc_obj_path == output_file_path):
            print("Warning: Input and Output files are same!")
        print("Updated file saved as: " + output_file_path)


def main():
    if(len(sys.argv) < 3):
        print("Not enough arguments!")
    elif(os.path.isfile(sys.argv[1]) == 0):
        print("No such file : " + sys.argv[1])
    else:
        replace_char = '*'
        input_file = sys.argv[1]
        regexes = [r"""\d{3}-\d{2}-\d{4}""", r"""(([a-zA-Z0-9_\.+-]+)@([a-zA-Z0-9-]+)\.[a-zA-Z0-9-\.]+)"""]
        redactor = DocxRedactor(input_file,regexes,replace_char)
        redactor.redact(sys.argv[2])

if __name__ == "__main__":
    main()
