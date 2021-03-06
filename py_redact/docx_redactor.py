import re

from docx import Document


class DocxRedactor:
    def __init__(self, doc_obj_path, regexes, replace_char):
        self.doc_obj_path = doc_obj_path
        self.regexes = regexes
        self.replace_char = replace_char

    def __redact_helper__(self, doc_obj):
        """
        Helper function for the redact function
        :param doc_obj (Document): the whole document as in the input .docx file
        :return:
        """
        for reg in self.regexes:
            regex = re.compile(reg)
            for section in doc_obj.sections:
                footer = section.footer
                header = section.header
                for paragraph in footer.paragraphs:
                    if regex.search(paragraph.text):
                        inline = paragraph.runs
                        for i in range(len(inline)):
                            match = regex.search(inline[i].text)
                            if match:
                                text = regex.sub(self.replace_char * len(match.group(0)), inline[i].text)
                                inline[i].text = text
                for paragraph in header.paragraphs:
                    if regex.search(paragraph.text):
                        inline = paragraph.runs
                        for i in range(len(inline)):
                            match = regex.search(inline[i].text)
                            if match:
                                text = regex.sub(self.replace_char * len(match.group(0)), inline[i].text)
                                inline[i].text = text
            for p in doc_obj.paragraphs:
                if regex.search(p.text):
                    inline = p.runs
                    for i in range(len(inline)):
                        match_obj = regex.search(inline[i].text)
                        if match_obj:
                            text = regex.sub(self.replace_char * len(match_obj.group(0)), inline[i].text)
                            inline[i].text = text
            for table in doc_obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self.__redact_helper__(cell)

    def redact(self, output_file_path):
        """
        Redacts the given .docx file and writes result to output file
        :param output_file_path (string): path of the file to write the result to
        :return:
        """
        doc_obj = Document(self.doc_obj_path)
        self.__redact_helper__(doc_obj)
        doc_obj.save(output_file_path)
        if self.doc_obj_path == output_file_path:
            print("Input and Output files are same!")
        print("Updated file saved as: " + output_file_path)
