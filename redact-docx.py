import sys
import os.path
import re
from docx import Document

def redact_docx(doc_obj, regexes, replace):
    for reg in regexes:
        regex = re.compile(reg)
        for p in doc_obj.paragraphs:
            if regex.search(p.text):
                inline = p.runs
                for i in range(len(inline)):
                    if regex.search(inline[i].text):
                        text = regex.sub(replace*len(inline[i].text), inline[i].text)
                        inline[i].text = text
                        
        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    redact_docx(cell, regexes, replace)
    return sys.argv[2]
    pass

if(len(sys.argv) < 3):
    print("Not enough arguments!")
elif(os.path.isfile(sys.argv[1]) == 0):
    print("No such file : " + sys.argv[1])
else:
    replace = '*'
    filename = sys.argv[1]
    doc_obj = Document(filename)
    regexes = [r"\d{3}-\d{2}-\d{4}", r"\d{3}-\d{3}-\d{4}"]
    Output_file = redact_docx(doc_obj, regexes, replace)
    doc_obj.save(Output_file)
    print("Updated file saved as : " + Output_file)